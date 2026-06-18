"""
app.py — Flask backend for Sofia, the career & business document platform.

Routes:
    POST /extract-text      PDF/Word -> plain text (no AI, saves tokens)
    POST /analyse-cv        Stage 1 analysis (Haiku)
    POST /rewrite-cv        Stage 2 rewrite (Sonnet)
    POST /cover-letter      Stage 3 cover letter (Sonnet)
    POST /rank-cvs          Engine 2 recruiter ranking (Sonnet)
    POST /generate-plan     Engine 3 business plan (Sonnet)
    POST /generate-pdfs     ReportLab PDFs (no AI)
    POST /generate-docx     python-docx Word files (no AI)
    GET  /health

Run:
    pip install flask flask-cors reportlab python-docx pdfplumber anthropic python-dotenv
    python app.py

Environment (.env):
    ANTHROPIC_API_KEY=sk-ant-...        (required)
    SOFIA_ALLOWED_ORIGINS=http://localhost:3000   (comma-separated; default '*')
    SOFIA_DEBUG=0                       (1 enables Flask debug — never in production)
    SOFIA_MAX_UPLOAD_MB=8               (max upload size)
"""

import base64
import io
import json
import logging
import os
import re
import time
import uuid
from collections import defaultdict
from functools import wraps
from logging.handlers import RotatingFileHandler
from threading import Lock

import anthropic
import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from flask import Flask, request, jsonify, g
from flask_cors import CORS

from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import HexColor
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

load_dotenv()

# --------------------------------------------------------------------------- #
#  LOGGING                                                                     #
# --------------------------------------------------------------------------- #
logging.basicConfig(level=logging.INFO)
log = logging.getLogger("sofia")


class _JsonFormatter(logging.Formatter):
    """One JSON object per line — easy to tail, parse, and forward."""
    def format(self, record):
        doc = {
            "ts": self.formatTime(record, "%Y-%m-%dT%H:%M:%S"),
            "level": record.levelname,
            "msg": record.getMessage(),
        }
        if record.exc_info:
            doc["exc"] = self.formatException(record.exc_info)
        return json.dumps(doc)


_fh = RotatingFileHandler("sofia.log", maxBytes=5 * 1024 * 1024, backupCount=3)
_fh.setLevel(logging.WARNING)
_fh.setFormatter(_JsonFormatter())
logging.getLogger().addHandler(_fh)

# Optional Sentry — activates when SENTRY_DSN env var is set.
# To enable: pip install sentry-sdk[flask]  and set SENTRY_DSN=https://...@sentry.io/...
try:
    import sentry_sdk
    from sentry_sdk.integrations.flask import FlaskIntegration
    from sentry_sdk.integrations.logging import LoggingIntegration as _SentryLogging

    _SENTRY_DSN = os.environ.get("SENTRY_DSN", "").strip()
    if _SENTRY_DSN:
        sentry_sdk.init(
            dsn=_SENTRY_DSN,
            integrations=[
                FlaskIntegration(),
                _SentryLogging(level=logging.WARNING, event_level=logging.ERROR),
            ],
            traces_sample_rate=0.1,
            environment=os.environ.get("SOFIA_ENV", "development"),
        )
        log.info("Sentry active (env=%s)", os.environ.get("SOFIA_ENV", "development"))
except ImportError:
    pass  # sentry-sdk not installed; pip install sentry-sdk[flask] when deploying

# --------------------------------------------------------------------------- #
#  CONFIG                                                                      #
# --------------------------------------------------------------------------- #
DEBUG_MODE     = os.environ.get("SOFIA_DEBUG", "0") == "1"
ALLOWED_ORIGINS = [o.strip() for o in os.environ.get("SOFIA_ALLOWED_ORIGINS", "*").split(",") if o.strip()]
MAX_UPLOAD_MB  = int(os.environ.get("SOFIA_MAX_UPLOAD_MB", "8"))
API_KEY        = os.environ.get("ANTHROPIC_API_KEY", "").strip()

# Hard input caps (characters) to bound token cost and block abuse (A6).
MAX_CV_CHARS  = 30_000
MAX_JD_CHARS  = 20_000
MAX_RANK_CVS  = 20
MAX_RANK_CHARS = 8_000   # per CV before the AI sees it

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_MB * 1024 * 1024
CORS(app, origins=ALLOWED_ORIGINS if ALLOWED_ORIGINS != ["*"] else "*")

SOFIA_ENV = os.environ.get("SOFIA_ENV", "development").strip()

# Production safety guard — refuse to start with unsafe defaults.
if SOFIA_ENV == "production":
    if DEBUG_MODE:
        raise SystemExit("FATAL: SOFIA_DEBUG=1 is not allowed in production.")
    if ALLOWED_ORIGINS == ["*"]:
        raise SystemExit("FATAL: SOFIA_ALLOWED_ORIGINS=* is not allowed in production. "
                         "Set it to your frontend domain.")
    if not API_KEY:
        raise SystemExit("FATAL: ANTHROPIC_API_KEY is not set in production.")

if not API_KEY:
    # Fail loud rather than booting a server whose every AI route 500s (A2).
    log.warning("ANTHROPIC_API_KEY is not set. AI routes will be unavailable until it is.")

_anthropic = anthropic.Anthropic(api_key=API_KEY) if API_KEY else None

HAIKU_MODEL  = "claude-haiku-4-5"
SONNET_MODEL = "claude-sonnet-4-6"

# --------------------------------------------------------------------------- #
#  CREDIT / AUTH SCAFFOLDING  (ADDITION — see audit note)                      #
#  The README documents a subscription + credit model that the original code  #
#  did not enforce. This is a pluggable, stubbed layer: replace `lookup_user`  #
#  and `charge_credits` with a real datastore. By default it is permissive so  #
#  local development is unchanged, but the hooks exist so launch can enforce.  #
# --------------------------------------------------------------------------- #
ENFORCE_CREDITS = os.environ.get("SOFIA_ENFORCE_CREDITS", "0") == "1"

ROUTE_COST = {            # credits per call, per README
    "/analyse-cv":    0,  # charged together with rewrite as a 2-credit action
    "/rewrite-cv":    2,
    "/cover-letter":  1,
    "/generate-plan": 5,
    "/rank-cvs":      1,  # per CV; multiplied by count at call site
}


class CreditError(Exception):
    def __init__(self, message, status=402):
        super().__init__(message)
        self.status = status


def lookup_user(token):
    """STUB. Return a user dict or None. Replace with a real lookup."""
    # Example shape: {"id": "...", "tier": "pro", "credits": 100, "unlimited": False}
    return None


def charge_credits(user, amount):
    """STUB. Deduct credits. Replace with an atomic datastore update."""
    if user is None or user.get("unlimited"):
        return
    if user.get("credits", 0) < amount:
        raise CreditError("Insufficient credits for this action.")
    user["credits"] -= amount


def require_credits(amount):
    """
    Returns (user, charge_fn). When enforcement is off, returns (None, noop)
    so behaviour is identical to the original permissive build.
    """
    if not ENFORCE_CREDITS:
        return None, (lambda: None)
    token = request.headers.get("Authorization", "").replace("Bearer ", "").strip()
    user = lookup_user(token)
    if user is None:
        raise CreditError("Authentication required.", status=401)
    if not user.get("unlimited") and user.get("credits", 0) < amount:
        raise CreditError("Insufficient credits for this action.")
    return user, (lambda: charge_credits(user, amount))


def user_is_paid():
    """True if the caller may see gated business-plan sections (A19)."""
    if not ENFORCE_CREDITS:
        return True   # local/dev: show everything
    token = request.headers.get("Authorization", "").replace("Bearer ", "").strip()
    user = lookup_user(token)
    return bool(user and (user.get("unlimited") or user.get("tier") in {"starter", "pro", "business"}))


# --------------------------------------------------------------------------- #
#  ERROR HELPERS  (A4 — never leak tracebacks)                                 #
# --------------------------------------------------------------------------- #
def fail(message, status=500, exc=None):
    err_id = uuid.uuid4().hex[:12]
    if exc is not None:
        log.exception("error_id=%s %s", err_id, message)
    else:
        log.error("error_id=%s %s", err_id, message)
    return jsonify({"status": "error", "message": message, "errorId": err_id}), status


@app.errorhandler(413)
def too_large(_e):
    return jsonify({"status": "error",
                    "message": f"File exceeds the {MAX_UPLOAD_MB} MB limit."}), 413


# --------------------------------------------------------------------------- #
#  RATE LIMITING (sliding window per IP — no external dependency)             #
#  Switch key from request.remote_addr to user ID once auth exists.           #
# --------------------------------------------------------------------------- #
_rl_store: dict = defaultdict(list)
_rl_lock = Lock()


def rate_limit(calls: int, window: int):
    """Allow at most `calls` requests per `window` seconds per client IP."""
    def decorator(f):
        @wraps(f)
        def wrapper(*args, **kwargs):
            ip = request.remote_addr or "unknown"
            key = f"{f.__name__}:{ip}"
            now = time.time()
            cutoff = now - window
            with _rl_lock:
                timestamps = [t for t in _rl_store[key] if t > cutoff]
                if len(timestamps) >= calls:
                    err_id = uuid.uuid4().hex[:12]
                    log.warning("rate_limit_exceeded error_id=%s route=%s ip=%s",
                                err_id, f.__name__, ip)
                    return jsonify({
                        "status": "error",
                        "message": "Too many requests. Please wait a moment and try again.",
                        "errorId": err_id,
                    }), 429
                timestamps.append(now)
                _rl_store[key] = timestamps
            return f(*args, **kwargs)
        return wrapper
    return decorator


@app.errorhandler(429)
def rate_limited(_e):
    return jsonify({"status": "error",
                    "message": "Too many requests. Please wait a moment and try again."}), 429


# --------------------------------------------------------------------------- #
#  FONTS                                                                       #
# --------------------------------------------------------------------------- #
FONT_DIR = "/usr/share/fonts/truetype/google-fonts/"
FONT_MAP = {
    "Sans":         "Poppins-Regular.ttf",
    "Sans-Bold":    "Poppins-Bold.ttf",
    "Sans-Italic":  "Poppins-Italic.ttf",
    "Sans-Med":     "Poppins-Medium.ttf",
    "Sans-Light":   "Poppins-Light.ttf",
    "Serif":        "Lora-Variable.ttf",
    "Serif-Italic": "Lora-Italic-Variable.ttf",
}
FALLBACK = {
    "Sans": "Helvetica", "Sans-Bold": "Helvetica-Bold",
    "Sans-Italic": "Helvetica-Oblique", "Sans-Med": "Helvetica-Bold",
    "Sans-Light": "Helvetica", "Serif": "Times-Roman",
    "Serif-Italic": "Times-Italic",
}
_REGISTERED = {}


def register_fonts():
    for alias, filename in FONT_MAP.items():
        try:
            pdfmetrics.registerFont(TTFont(alias, FONT_DIR + filename))
            _REGISTERED[alias] = alias
        except Exception:
            _REGISTERED[alias] = FALLBACK[alias]


def F(alias):
    return _REGISTERED.get(alias, FALLBACK.get(alias, "Helvetica"))


register_fonts()

# --------------------------------------------------------------------------- #
#  COLOURS                                                                     #
# --------------------------------------------------------------------------- #
NIGHT = HexColor("#111111")
SLATE = HexColor("#333333")
MID   = HexColor("#777777")
LIGHT = HexColor("#BBBBBB")
PAPER = HexColor("#F7F5F1")
GOLD  = HexColor("#B89A68")
WHITE = HexColor("#FFFFFF")

PAGE_W, PAGE_H = A4

# --------------------------------------------------------------------------- #
#  PROMPTS                                                                     #
#  Static shape/instruction text lives in the SYSTEM block so prompt caching   #
#  actually saves tokens on repeat calls (A15). Only variable CV/JD goes in    #
#  the user message.                                                           #
# --------------------------------------------------------------------------- #
STAGE1_SHAPE = """{
  "firstImpression": { "perceivedRole": str, "perceivedSeniority": str, "strongestSignal": str, "biggestRedFlag": str, "verdict": "Shortlist"|"Maybe"|"Discard", "verdictReason": str },
  "scorecard": {
    "contentQuality": [ { "category": str, "score": num, "issue": str, "quote": str } ],
    "strategicFit": [ { "category": str, "score": num, "issue": str, "quote": str } ],
    "presentationTrust": [ { "category": str, "score": num, "issue": str, "quote": str } ],
    "overallScore": num, "grade": "A"|"B"|"C"|"D", "gradeLabel": str
  },
  "killList": [ { "weakLine": str, "whyItFails": str, "replacement": str } ],
  "competitiveIntelligence": { "topShortlistReasons": [str,str,str], "interviewWeaknesses": [str,str], "competitionAssessment": str, "quickWins": [ { "action": str, "timeEstimate": str, "whereToShow": str } ], "honestLimitation": str }
}
Content quality (8 categories): Impact & Quantification, Action Verb Strength, Bullet Quality, Bullet Discipline, Summary Power, Skills Section, Verb Tense Consistency, Honesty Signal.
Strategic fit (4): ATS Keyword Match, Role Alignment, Seniority Match, Competitive Edge.
Presentation & trust (4): Formatting & ATS Safety, Readability, Career Narrative, Gap/Regression Handling."""

STAGE1_SYSTEM = """You are a Senior Talent Partner with 15 years recruiting at top-tier companies.
You think like a hiring manager scanning 200 CVs — not a keyword tool.
Evaluate this CV with brutal precision. Generic feedback is failure.

SCORING GUIDE: 9-10 Near perfect. 7-8 Good, minor fixes. 5-6 Average.
3-4 Weak, actively hurting chances. 1-2 Damaging, must be removed.
Grade (out of 160): A 136-160. B 112-135. C 80-111. D below 80.

Quote weak lines EXACTLY. Write replacements exactly as they should appear.
Never fabricate metrics. Flag missing ones as [ADD METRIC].
Never use "estimate", "est.", "approximately", or "roughly" — use hard numbers only.
Return ONLY valid JSON matching this exact shape. No markdown. No preamble:
""" + STAGE1_SHAPE

STAGE1_NO_JD_ADDENDUM = """

No job description will be provided. Evaluate for general professional positioning strength.
Replace all JD-matching criteria with: Does this CV position the candidate as a strong professional in their apparent field?"""

STAGE2_SHAPE = """{
  "improvementPlan": {
    "priorityFixes": [ { "section": str, "problem": str, "fix": str, "impact": "High"|"Medium"|"Low" } ],
    "keywordsToAdd": [str], "keywordsToRemove": [str],
    "summaryRewrite": str,
    "bulletTransformations": [ { "original": str, "rewritten": str, "reason": str } ]
  },
  "rewrittenCV": {
    "fullText": str, "header": str, "summary": str, "experience": str, "skills": str, "education": str, "certifications": str,
    "candidateName": str, "candidateEmail": str, "candidatePhone": str, "candidateLinkedIn": str, "candidateWebsite": str, "candidateLocation": str,
    "titleLine": str, "companyName": str
  },
  "pdfData": {
    "expertise": [str], "tools": [str],
    "education": [ { "degree": str, "school": str, "years": str } ],
    "certifications": [ { "name": str, "org": str } ],
    "projects": [str],
    "experienceItems": [ { "role": str, "company": str, "location": str, "dates": str, "bullets": [str] } ]
  },
  "rescore": {
    "categories": [ { "category": str, "before": num, "after": num } ],
    "scoreBefore": num, "scoreAfter": num, "gradeBefore": str, "gradeAfter": str
  },
  "recruiterNote": {
    "shortlistReasons": [str,str,str], "interviewWeaknesses": [str,str],
    "quickWins": [ { "action": str, "timeEstimate": str, "whereToShow": str } ]
  },
  "interviewPrep": {
    "questions": [ { "question": str, "weakness": str, "starAnswer": { "situation": str, "task": str, "action": str, "result": str }, "needsRealExample": bool } ]
  }
}"""

STAGE2_SYSTEM = """You are the same Senior Talent Partner. Rewrite the CV and build the full improvement package.

REWRITE RULES (no exceptions):
- Every bullet contains a metric. Unknown: [ADD METRIC]. Never "estimate" or "approximately".
- Every bullet opens with a unique strong past-tense action verb. No repeated verbs per role.
- Every bullet is outcome-focused. No tasks without results.
- 3-5 bullets per role. Hard limit.
- Summary: seniority + years + value proposition. Max 4 lines.
- Hard skills only in skills section. Zero soft skills.
- No pronouns. No buzzwords. No filler.
- ATS-safe: no tables, no columns, no text boxes.
- Never fabricate. Reframe only. Flag unknowns as [ADD METRIC] or [VERIFY CLAIM].

STAR answers must reference real CV achievements. Flag missing examples as [NEEDS REAL EXAMPLE].
Return ONLY valid JSON matching this exact shape. No markdown. No preamble:
""" + STAGE2_SHAPE

STAGE2_NO_JD_ADDENDUM = """

No job description will be provided. Strengthen for general professional positioning.
Remove all JD mirroring instructions. Focus on universal CV strength principles.
Leave keywordsToAdd and keywordsToRemove as empty arrays."""

STAGE3_SHAPE = """{ "coverLetter": { "opening": str, "body1": str, "body2": str, "body3Remote": str, "closing": str, "fullText": str, "signoffName": str, "tagline": str } }"""

STAGE3_SYSTEM = """You write executive-level cover letters that sound human, sharp, and never AI-generated.
No headings in the letter. Max 500 words. Every sentence earns its place.
Never use "estimate", "est.", "approximately", or "roughly".

Paragraph structure:
- opening: Role, why applying, three clear reasons of fit.
- body1 (Impact): 2-3 specific CV achievements linked to company needs with metrics.
- body2 (Company interest): Product, mission, growth stage — researched, genuine.
- body3Remote (natural prose, not a list): timezone compatibility with the company location; if and only if the candidate's location is given and differs from the company's, ONE matter-of-fact sentence confirming international payment is set up — never invent a country; first 90 days — 1-2 concrete things to build or fix.
- closing: Reaffirm interest. Reinforce value. Invite conversation. Confident, not desperate.
Return ONLY valid JSON matching this exact shape. No markdown. No preamble:
""" + STAGE3_SHAPE

RANK_CVS_SHAPE = """{
  "rankings": [
    { "candidateIndex": num, "rank": num, "score": num, "topStrengths": [str,str], "keyGaps": [str], "hiringRecommendation": str }
  ],
  "top3Summary": str,
  "commonWeaknesses": [str]
}"""

RANK_CVS_SYSTEM = """You are a senior recruiter reviewing multiple CVs for a single role.
Rank ALL candidates objectively (every candidate index must appear exactly once).
Give concrete reasons for the top 3. Never use "estimate" or vague language.
Return ONLY valid JSON matching this exact shape. No markdown. No preamble:
""" + RANK_CVS_SHAPE

PLAN_SYSTEM = """You are an expert business plan writer with deep experience in African markets and grant applications.
Write business plans that are human, confident, and grant-appropriate.
All numbers must be internally consistent. The use-of-funds items must add up EXACTLY to the total ask.
Never use "estimate", "est.", "approximately", or "roughly" — use specific projections.
Return ONLY valid JSON. No markdown. No preamble."""


# --------------------------------------------------------------------------- #
#  GRANT LIBRARY                                                              #
# --------------------------------------------------------------------------- #
GRANT_LIBRARY = {
    "tef": {
        "name": "Tony Elumelu Foundation Entrepreneurship Programme",
        "required_sections": ["executiveSummary", "problemStatement", "solution", "marketOpportunity", "revenueModel", "impactProjection", "useOfFunds"],
        "free_sections": ["executiveSummary", "marketOpportunity"],
        "usp_guidance": "Lead with African job creation numbers. Quantify community impact. Show founder-market fit. Emphasise scalability across the continent. TEF prioritises businesses that create jobs and transform African communities.",
        "tone": "Passionate, impact-driven, community-focused. Write as an African entrepreneur who deeply understands their local market.",
    },
    "hult": {
        "name": "Hult Prize",
        "required_sections": ["executiveSummary", "problemAndSDG", "solution", "businessModel", "marketSize", "socialImpactMetrics", "team", "financialAsk"],
        "free_sections": ["executiveSummary", "problemAndSDG"],
        "usp_guidance": "Open with the specific UN SDG addressed. Use global scale language. Impact metrics must be specific and measurable. Judges are MBA-educated — financial model rigour is non-negotiable.",
        "tone": "Rigorous, data-led, globally minded. Academic in structure but passionate in mission.",
    },
    "yali": {
        "name": "YALI Mandela Washington Fellowship — Business Track",
        "required_sections": ["executiveSummary", "leadershipStory", "businessOverview", "communityImpact", "sustainabilityPlan", "useOfFunds"],
        "free_sections": ["executiveSummary", "businessOverview"],
        "usp_guidance": "Centre the founder's personal leadership journey. Connect business goals to community upliftment. Show how U.S. exchange creates lasting value back home. YALI values character as much as business model.",
        "tone": "Personal, leadership-focused, values-driven. Write in first person for narrative sections.",
    },
    "seedstars": {
        "name": "Seedstars Africa Ventures",
        "required_sections": ["executiveSummary", "problem", "solutionAndProduct", "marketAnalysis", "tractionAndMetrics", "businessModel", "financials", "team", "theAsk"],
        "free_sections": ["executiveSummary", "marketAnalysis"],
        "usp_guidance": "Lead with traction — month-over-month growth, active users, revenue if any. Investors want numbers. TAM must be Africa-sized, not global. Unit economics must be defensible. Seedstars invests, not grants — show path to profitability.",
        "tone": "Investor-grade. Concise, metrics-first, growth-obsessed.",
    },
    "general": {
        "name": "General Business Plan",
        "required_sections": ["executiveSummary", "companyOverview", "marketAnalysis", "productsAndServices", "marketingStrategy", "operationsPlan", "financialProjections", "managementTeam", "useOfFunds"],
        "free_sections": ["executiveSummary", "marketAnalysis"],
        "usp_guidance": "Tailor tone to the reader. For investors: growth and return. For banks: cash flow and collateral. For internal: operations and milestones.",
        "tone": "Professional, clear, confident.",
    },
}


def free_sections_for(grant):
    """Per-grant free preview (A12). Falls back to the first two required sections."""
    fs = grant.get("free_sections")
    if fs:
        return list(fs)
    return list(grant["required_sections"][:2])


# --------------------------------------------------------------------------- #
#  RULES ENGINE (no AI cost)                                                  #
# --------------------------------------------------------------------------- #
GENERIC_PHRASES = {
    "proven track record", "results-driven", "dynamic professional", "team player",
    "self-starter", "go-getter", "hard worker", "passionate about", "detail-oriented",
    "excellent communication", "strong interpersonal", "highly motivated", "fast learner",
    "problem solver", "think outside the box", "synergy", "leverage", "value-add",
    "best practices", "hit the ground running", "proactive approach"
}

STRONG_VERBS = {
    "achieved", "accelerated", "built", "closed", "created", "delivered", "designed",
    "developed", "drove", "engineered", "executed", "expanded", "generated", "grew",
    "implemented", "increased", "launched", "led", "managed", "optimised", "optimized",
    "produced", "reduced", "scaled", "secured", "shipped", "spearheaded", "transformed"
}

# Current-year-relative window for "has dates" detection (A10).
import datetime as _dt
_RECENT_YEARS = {str(y) for y in range(_dt.date.today().year - 12, _dt.date.today().year + 1)}


def _bullet_lines(cv_text):
    """
    Identify bullet-like lines. Extracted text often loses glyphs, so we also
    treat short standalone lines that begin with a strong verb as bullets (A10).
    """
    lines = [ln.strip() for ln in cv_text.split("\n") if ln.strip()]
    bullets = []
    for ln in lines:
        if ln[0] in "•-\u25aa*\u2022\u2023\u2043":
            bullets.append(ln.lstrip("•-\u25aa*\u2022\u2023\u2043 ").strip())
        else:
            first = ln.split()[0].lower().strip(".,") if ln.split() else ""
            if first in STRONG_VERBS and len(ln.split()) >= 4:
                bullets.append(ln)
    return bullets


def rules_engine(cv_text: str) -> dict:
    text_lower = cv_text.lower()
    words = cv_text.split()
    sentences = [s.strip() for s in re.split(r'[.!\n]', cv_text) if s.strip()]

    ats_deductions = 0
    if re.search(r'\|.*\|', cv_text): ats_deductions += 20
    if re.search(r'  {4,}', cv_text): ats_deductions += 10
    ats_score = max(0, 100 - ats_deductions)

    bullets = _bullet_lines(cv_text)
    metric_bullets = sum(1 for b in bullets if re.search(r'\d', b))
    metric_density = (metric_bullets / max(len(bullets), 1)) * 100

    first_words = [s.split()[0].lower() if s.split() else '' for s in sentences]
    strong_count = sum(1 for w in first_words if w in STRONG_VERBS)
    verb_score = min(100, (strong_count / max(len(sentences), 1)) * 200)

    generic_count = sum(1 for phrase in GENERIC_PHRASES if phrase in text_lower)
    generic_density = (generic_count / max(len(words) / 50, 1))

    action_verbs_used = [w for w in first_words if w in STRONG_VERBS]
    verb_diversity = len(set(action_verbs_used)) / max(len(action_verbs_used), 1)

    bullet_lengths = [len(b.split()) for b in bullets]
    if len(bullet_lengths) > 3:
        avg = sum(bullet_lengths) / len(bullet_lengths)
        variance = sum((l - avg) ** 2 for l in bullet_lengths) / len(bullet_lengths)
        high_uniformity = variance < 10
    else:
        high_uniformity = False

    has_dates = any(y in cv_text for y in _RECENT_YEARS)
    ai_flags = sum([
        verb_diversity < 0.4,
        generic_density > 0.15,
        metric_density < 10,
        high_uniformity,
        not has_dates,
    ])
    ai_flag = ai_flags >= 3

    expected_sections = ["experience", "education", "skills", "summary", "contact"]
    found_sections = [s for s in expected_sections if s in text_lower]

    return {
        "ats_score": ats_score,
        "verb_strength": round(verb_score),
        "metric_density": round(metric_density),
        "generic_density": round(generic_density * 100),
        "ai_flag": ai_flag,
        "word_count": len(words),
        "section_completeness": {
            "found": found_sections,
            "missing": [s for s in expected_sections if s not in found_sections],
        },
    }


# --------------------------------------------------------------------------- #
#  AI CALL HELPER                                                             #
# --------------------------------------------------------------------------- #
def call_claude(model: str, system: str, user: str, max_tokens: int):
    """
    Call Anthropic with prompt caching on the (now large, static) system prompt.
    Returns (text, stop_reason). Raises RuntimeError if the client is unconfigured.
    """
    if _anthropic is None:
        raise RuntimeError("AI is unavailable: ANTHROPIC_API_KEY is not configured.")
    response = _anthropic.messages.create(
        model=model,
        max_tokens=max_tokens,
        system=[{"type": "text", "text": system, "cache_control": {"type": "ephemeral"}}],
        messages=[{"role": "user", "content": user}],
        timeout=120,
    )
    text = "".join(b.text for b in response.content if b.type == "text")
    return text, response.stop_reason


def parse_json(raw: str) -> dict:
    """Tolerant JSON extraction: strip fences, take the outermost object."""
    s = (raw or "").strip()
    s = re.sub(r'^```(?:json)?', '', s, flags=re.IGNORECASE).strip()
    s = re.sub(r'```$', '', s).strip()
    a, b = s.find('{'), s.rfind('}')
    if a != -1 and b != -1 and b > a:
        s = s[a:b + 1]
    return json.loads(s)


def ai_json(model, system, user, max_tokens):
    """
    Call the model and parse JSON, with one retry on malformed/truncated output (A8/A9).
    Raises ValueError with a clean message on persistent failure.
    """
    last_err = None
    for attempt in range(2):
        text, stop = call_claude(model, system, user, max_tokens)
        if stop == "max_tokens":
            last_err = "The model response was cut off (max_tokens). Try a shorter input."
            # Still attempt a parse — partial objects sometimes parse after trimming.
        try:
            return parse_json(text)
        except (json.JSONDecodeError, ValueError) as e:
            last_err = "The model returned output that was not valid JSON."
            log.warning("JSON parse failed (attempt %d): %s", attempt + 1, e)
    raise ValueError(last_err or "The model returned malformed output.")


def compress_stage1(s1: dict) -> dict:
    """Extract only what Stage 2 needs — cuts ~60% of tokens."""
    sc = s1.get("scorecard", {})
    ci = s1.get("competitiveIntelligence", {})
    return {
        "overallScore": sc.get("overallScore"),
        "grade": sc.get("grade"),
        "killList": s1.get("killList", [])[:5],
        "interviewWeaknesses": ci.get("interviewWeaknesses", []),
        "priorityFixes": [
            f for f in (
                sc.get("contentQuality", []) +
                sc.get("strategicFit", []) +
                sc.get("presentationTrust", [])
            ) if f.get("score", 10) < 7
        ][:8],
    }


# --------------------------------------------------------------------------- #
#  TEXT EXTRACTION                                                            #
# --------------------------------------------------------------------------- #
def extract_pdf_text(file_bytes: bytes) -> str:
    lines = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if page_text:
                lines.append(page_text)
    raw = "\n".join(lines)
    raw = re.sub(r'[^\x20-\x7E\n]', ' ', raw)
    raw = re.sub(r' {3,}', '  ', raw)
    raw = re.sub(r'\n{3,}', '\n\n', raw)
    return raw.strip()


def extract_docx_text(file_bytes: bytes) -> str:
    from docx.oxml.ns import qn
    doc = Document(io.BytesIO(file_bytes))
    lines = []

    def _collect_para(para):
        text = para.text.strip()
        if not text:
            return
        if para.style.name.startswith('Heading'):
            lines.append(text.upper())
        else:
            lines.append(text)

    # Main body text flow
    for para in doc.paragraphs:
        _collect_para(para)

    # Table cells — not included in doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _collect_para(para)

    # Text boxes — floating frames invisible to the python-docx API
    for shape in doc.element.body.iter(qn('w:txbxContent')):
        for p_elem in shape.iter(qn('w:p')):
            text = "".join(t.text for t in p_elem.iter(qn('w:t'))).strip()
            if text:
                lines.append(text)

    raw = "\n".join(lines)
    raw = re.sub(r'\n{3,}', '\n\n', raw)
    return raw.strip()


def detect_sections(text: str) -> list:
    common = ["experience", "education", "skills", "summary", "certifications",
              "projects", "awards", "languages", "publications", "contact", "profile"]
    found = []
    for line in text.split("\n"):
        stripped = line.strip().lower()
        if stripped in common:
            found.append(line.strip().upper())
    return list(dict.fromkeys(found))


# --------------------------------------------------------------------------- #
#  ROUTE — TEXT EXTRACTION                                                    #
# --------------------------------------------------------------------------- #
@app.route("/extract-text", methods=["POST"])
@rate_limit(30, 3600)   # 30 per hour — prevent free file-parsing abuse
def extract_text():
    try:
        if "file" not in request.files:
            return jsonify({"status": "error", "message": "No file uploaded"}), 400
        f = request.files["file"]
        file_bytes = f.read()
        filename = (f.filename or "").lower()

        if filename.endswith(".pdf"):
            text = extract_pdf_text(file_bytes)
        elif filename.endswith(".docx"):
            text = extract_docx_text(file_bytes)
        elif filename.endswith(".doc"):
            return jsonify({"status": "error",
                            "message": "Old .doc format is not supported. Re-save your CV as .docx in Word and try again."}), 400
        else:
            return jsonify({"status": "error",
                            "message": "Unsupported file type. Upload a PDF or Word (.docx) file."}), 400

        if len(text) > MAX_CV_CHARS:
            text = text[:MAX_CV_CHARS]

        return jsonify({
            "status": "success",
            "text": text,
            "sections": detect_sections(text),
            "wordCount": len(text.split()),
            "rulesEngine": rules_engine(text),
        })
    except Exception as e:
        return fail("Could not read that file. It may be corrupt or password-protected.",
                    500, exc=e)


# --------------------------------------------------------------------------- #
#  ROUTES — AI CALLS                                                          #
# --------------------------------------------------------------------------- #
@app.route("/analyse-cv", methods=["POST"])
@rate_limit(10, 3600)   # 10 per hour
def analyse_cv():
    try:
        require_credits(ROUTE_COST["/analyse-cv"])  # 0 by design; pairs with rewrite
        data = request.get_json(force=True) or {}
        cv_text = (data.get("cvText") or "").strip()[:MAX_CV_CHARS]
        jd_text = (data.get("jdText") or "").strip()[:MAX_JD_CHARS]
        if not cv_text:
            return jsonify({"status": "error", "message": "cvText is required"}), 400

        system = STAGE1_SYSTEM + ("" if jd_text else STAGE1_NO_JD_ADDENDUM)
        user = f"=== CV ===\n{cv_text}"
        if jd_text:
            user += f"\n\n=== JOB DESCRIPTION ===\n{jd_text}"

        result = ai_json(HAIKU_MODEL, system, user, max_tokens=2500)
        return jsonify({"status": "success", "data": result})
    except CreditError as ce:
        return fail(str(ce), ce.status)
    except ValueError as ve:
        return fail(str(ve), 502)
    except Exception as e:
        return fail("Analysis failed. Please try again.", 500, exc=e)


@app.route("/rewrite-cv", methods=["POST"])
@rate_limit(10, 3600)   # 10 per hour
def rewrite_cv():
    try:
        user_obj, charge = require_credits(ROUTE_COST["/rewrite-cv"])
        data = request.get_json(force=True) or {}
        cv_text = (data.get("cvText") or "").strip()[:MAX_CV_CHARS]
        jd_text = (data.get("jdText") or "").strip()[:MAX_JD_CHARS]
        stage1_summary = data.get("stage1Summary") or {}
        if not cv_text:
            return jsonify({"status": "error", "message": "cvText is required"}), 400

        system = STAGE2_SYSTEM + ("" if jd_text else STAGE2_NO_JD_ADDENDUM)
        user = (
            f"=== STAGE 1 SUMMARY ===\n{json.dumps(stage1_summary)}\n\n"
            f"=== ORIGINAL CV ===\n{cv_text}"
        )
        if jd_text:
            user += f"\n\n=== JOB DESCRIPTION ===\n{jd_text}"

        result = ai_json(SONNET_MODEL, system, user, max_tokens=6000)
        charge()  # only charge once the call succeeded
        return jsonify({"status": "success", "data": result})
    except CreditError as ce:
        return fail(str(ce), ce.status)
    except ValueError as ve:
        return fail(str(ve), 502)
    except Exception as e:
        return fail("Rewrite failed. Please try again.", 500, exc=e)


@app.route("/cover-letter", methods=["POST"])
@rate_limit(10, 3600)   # 10 per hour
def cover_letter():
    try:
        _u, charge = require_credits(ROUTE_COST["/cover-letter"])
        data = request.get_json(force=True) or {}
        approved_cv = (data.get("approvedCV") or "").strip()[:MAX_CV_CHARS]
        jd_text = (data.get("jdText") or "").strip()[:MAX_JD_CHARS]
        candidate_meta = data.get("candidateMeta") or {}
        if not approved_cv:
            return jsonify({"status": "error", "message": "approvedCV is required"}), 400
        if not jd_text:
            return jsonify({"status": "error",
                            "message": "A job description is required for the cover letter."}), 400

        user = (
            f"=== APPROVED CV ===\n{approved_cv}\n\n"
            f"=== CANDIDATE ===\n{json.dumps(candidate_meta)}\n\n"
            f"=== JOB DESCRIPTION ===\n{jd_text}"
        )
        result = ai_json(SONNET_MODEL, STAGE3_SYSTEM, user, max_tokens=2000)
        charge()
        return jsonify({"status": "success", "data": result})
    except CreditError as ce:
        return fail(str(ce), ce.status)
    except ValueError as ve:
        return fail(str(ve), 502)
    except Exception as e:
        return fail("Cover letter generation failed. Please try again.", 500, exc=e)


@app.route("/rank-cvs", methods=["POST"])
@rate_limit(5, 3600)    # 5 per hour — each call can rank up to 20 CVs
def rank_cvs():
    try:
        data = request.get_json(force=True) or {}
        cv_texts = data.get("cvTexts") or []
        jd_text = (data.get("jdText") or "").strip()[:MAX_JD_CHARS]
        if not cv_texts:
            return jsonify({"status": "error", "message": "cvTexts array is required"}), 400
        if len(cv_texts) > MAX_RANK_CVS:
            return jsonify({"status": "error",
                            "message": f"Upload at most {MAX_RANK_CVS} CVs per batch."}), 400

        _u, charge = require_credits(ROUTE_COST["/rank-cvs"] * len(cv_texts))

        rules_results = [rules_engine(cv) for cv in cv_texts]
        cv_block = "\n\n".join(
            f"=== CANDIDATE {i+1} ===\n{cv[:MAX_RANK_CHARS]}"
            for i, cv in enumerate(cv_texts)
        )
        user = cv_block + (f"\n\n=== ROLE ===\n{jd_text}" if jd_text else "")
        result = ai_json(SONNET_MODEL, RANK_CVS_SYSTEM, user, max_tokens=3000)

        # Validate + merge rules data; ensure every candidate is covered (A14).
        rankings = result.get("rankings") or []
        seen = set()
        for item in rankings:
            idx = (item.get("candidateIndex") or 0) - 1
            if 0 <= idx < len(rules_results) and idx not in seen:
                item["rulesEngine"] = rules_results[idx]
                seen.add(idx)
        for i in range(len(cv_texts)):
            if i not in seen:
                rankings.append({
                    "candidateIndex": i + 1, "rank": None, "score": None,
                    "topStrengths": [], "keyGaps": ["Not ranked by the model — review manually."],
                    "hiringRecommendation": "Review", "rulesEngine": rules_results[i],
                })
        result["rankings"] = rankings

        charge()
        return jsonify({"status": "success", "data": result})
    except CreditError as ce:
        return fail(str(ce), ce.status)
    except ValueError as ve:
        return fail(str(ve), 502)
    except Exception as e:
        return fail("Ranking failed. Please try again.", 500, exc=e)


@app.route("/generate-plan", methods=["POST"])
@rate_limit(5, 3600)    # 5 per hour — 5 credits per call
def generate_plan():
    try:
        _u, charge = require_credits(ROUTE_COST["/generate-plan"])
        data = request.get_json(force=True) or {}
        grant_id = (data.get("grantId") or "general").lower()
        form = data.get("formData") or {}

        grant = GRANT_LIBRARY.get(grant_id, GRANT_LIBRARY["general"])
        sections = grant["required_sections"]
        free = free_sections_for(grant)

        # Structured financials so the "use of funds totals the ask" rule is checkable (A13).
        shape = {s: "str" for s in sections}
        shape["useOfFundsTable"] = [{"item": "str", "amountUSD": "num"}]
        shape_str = json.dumps(shape, indent=2)

        system = (
            PLAN_SYSTEM
            + f"\n\nGrant context: {grant['name']}\n"
            + f"USP guidance: {grant['usp_guidance']}\n"
            + f"Tone: {grant['tone']}"
        )
        funding = str(form.get("fundingAmount", "0"))
        user = (
            f"Return a JSON object with exactly these keys:\n{shape_str}\n\n"
            f"=== BUSINESS DETAILS ===\n"
            f"Business name: {form.get('businessName', '')}\n"
            f"Industry: {form.get('industry', '')}\n"
            f"Country of operation: {form.get('country', 'Nigeria')}\n"
            f"Business stage: {form.get('stage', '')}\n"
            f"Employees: {form.get('employees', '')}\n"
            f"Target funding (USD): ${funding}\n"
            f"Use of funds: {form.get('useOfFunds', 'Not specified')}\n"
            f"Brief description: {form.get('description', '')}\n\n"
            f"Write each section as flowing, grant-quality prose. "
            f"The useOfFundsTable amounts MUST sum to exactly {funding} USD. "
            f"Never use 'estimate' or 'approximately'."
        )
        result = ai_json(SONNET_MODEL, system, user, max_tokens=5000)

        # Server-side validation of the funding total.
        warnings = []
        try:
            total = sum(float(r.get("amountUSD") or 0) for r in result.get("useOfFundsTable", []))
            ask = float(re.sub(r"[^\d.]", "", funding) or 0)
            if ask and abs(total - ask) > 0.5:
                warnings.append(f"Use-of-funds total ({total:.0f}) does not match the ask ({ask:.0f}).")
        except (TypeError, ValueError):
            warnings.append("Could not validate the use-of-funds total.")

        # Gate paid content SERVER-SIDE so free users never receive it (A19).
        paid = user_is_paid()
        gated = {k: (k not in free) for k in result if k != "useOfFundsTable"}
        sent = dict(result)
        if not paid:
            for k, is_gated in gated.items():
                if is_gated and k in sent:
                    sent[k] = None  # withhold body; frontend shows an upgrade prompt
            if "useOfFundsTable" in sent and "useOfFunds" not in free:
                sent["useOfFundsTable"] = None

        charge()
        return jsonify({
            "status": "success",
            "data": sent,
            "gated": gated,
            "paid": paid,
            "grant": grant["name"],
            "freeSections": free,
            "warnings": warnings,
        })
    except CreditError as ce:
        return fail(str(ce), ce.status)
    except ValueError as ve:
        return fail(str(ve), 502)
    except Exception as e:
        return fail("Business plan generation failed. Please try again.", 500, exc=e)


# --------------------------------------------------------------------------- #
#  LOW-LEVEL PDF HELPERS                                                      #
# --------------------------------------------------------------------------- #
def _wrap_lines(c, text, max_w, font, size):
    text = (text or "").replace("\r", " ").strip()
    if not text:
        return []
    lines, current = [], ""
    for word in text.split():
        trial = word if not current else current + " " + word
        if c.stringWidth(trial, font, size) <= max_w:
            current = trial
        else:
            if current:
                lines.append(current)
            if c.stringWidth(word, font, size) > max_w:
                chunk = ""
                for ch in word:
                    if c.stringWidth(chunk + ch, font, size) <= max_w:
                        chunk += ch
                    else:
                        lines.append(chunk)
                        chunk = ch
                current = chunk
            else:
                current = word
    if current:
        lines.append(current)
    return lines


def wrap(c, text, x, y, max_w, font, size, color, lh=None):
    lh = lh if lh is not None else size * 1.5
    c.setFillColor(color)
    c.setFont(font, size)
    for line in _wrap_lines(c, text, max_w, font, size):
        c.drawString(x, y, line)
        y -= lh
    return y


def hr(c, x1, y, x2, color=LIGHT, lw=0.35):
    c.setStrokeColor(color)
    c.setLineWidth(lw)
    c.line(x1, y, x2, y)


def section_label(c, text, x, y, w):
    spaced = "  ".join(list(text.upper()))
    c.setFillColor(GOLD)
    c.setFont(F("Sans-Bold"), 6.2)
    c.drawString(x, y, spaced)
    hr(c, x, y - 5, x + w, color=GOLD, lw=0.6)
    return y - 16


def bullet_line(c, text, x, y, max_w, font="Serif", size=8.4, color=SLATE):
    c.setFillColor(GOLD)
    c.setFont(F("Sans-Bold"), 5.5)
    c.drawString(x, y + 1.5, "\u25aa")
    lh = size * 1.52
    c.setFillColor(color)
    c.setFont(F(font), size)
    for line in _wrap_lines(c, text, max_w - 10, F(font), size):
        c.drawString(x + 10, y, line)
        y -= lh
    return y - 3


def footer(c, name, email, website):
    c.setFillColor(GOLD)
    c.rect(0, 0, PAGE_W, 4, fill=1, stroke=0)
    parts = [p for p in [name, email, website] if p]
    line = "   \u00b7   ".join(parts)
    size = 6.2
    c.setFont(F("Sans"), size)
    c.setFillColor(MID)
    tw = c.stringWidth(line, F("Sans"), size)
    x = (PAGE_W - tw) / 2.0
    y = 14
    c.drawString(x, y, line)
    cursor = x
    sep_w = c.stringWidth("   \u00b7   ", F("Sans"), size)
    for i, p in enumerate(parts):
        seg_w = c.stringWidth(p, F("Sans"), size)
        if "@" in p and " " not in p:
            c.linkURL("mailto:" + p, (cursor, y - 2, cursor + seg_w, y + size), relative=0, thickness=0)
        elif "." in p and " " not in p:
            u = p if p.startswith("http") else "https://" + p
            c.linkURL(u, (cursor, y - 2, cursor + seg_w, y + size), relative=0, thickness=0)
        cursor += seg_w + (sep_w if i < len(parts) - 1 else 0)


# --------------------------------------------------------------------------- #
#  PDF — TEMPLATE A (sidebar)                                                 #
# --------------------------------------------------------------------------- #
CV_HEADER_H  = 106
CV_SIDEBAR_W = 162
CV_CONTENT_X = 188
CV_CONTENT_W = PAGE_W - CV_CONTENT_X - 28


def draw_bg_a(c):
    c.setFillColor(PAPER)
    c.rect(0, 0, CV_SIDEBAR_W, PAGE_H, fill=1, stroke=0)
    c.setFillColor(GOLD)
    c.rect(0, 0, 4, PAGE_H, fill=1, stroke=0)


def cv_header_a(c, name, title_line, contact_items):
    c.setFillColor(NIGHT)
    c.rect(0, PAGE_H - CV_HEADER_H, PAGE_W, CV_HEADER_H, fill=1, stroke=0)
    c.setFillColor(GOLD)
    c.rect(0, PAGE_H - 4, PAGE_W, 4, fill=1, stroke=0)
    name_size = min(22, max(14, int(220 / max(len(name), 1))))
    c.setFillColor(WHITE)
    c.setFont(F("Serif"), name_size)
    c.drawString(20, PAGE_H - 38, name)
    if title_line:
        c.setFillColor(GOLD)
        c.setFont(F("Sans-Med"), 8)
        c.drawString(20, PAGE_H - 56, title_line)
    contact_str = "   \u00b7   ".join(contact_items[:4])
    c.setFillColor(WHITE)
    c.setFont(F("Sans"), 7.5)
    c.drawString(20, PAGE_H - 78, contact_str)


def build_cv_pdf_a(cv: dict, pdf_data: dict) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)

    name = cv.get("candidateName") or "Candidate"
    email = cv.get("candidateEmail") or ""
    linkedin = cv.get("candidateLinkedIn") or ""
    website = cv.get("candidateWebsite") or ""
    phone = cv.get("candidatePhone") or ""
    title_line = cv.get("titleLine") or ""
    location = cv.get("candidateLocation") or ""
    summary_text = cv.get("summary") or ""
    contact_items = [i for i in [email, phone, linkedin, location] if i]

    def new_page():
        draw_bg_a(c)
        cv_header_a(c, name, title_line, contact_items)
        footer(c, name, email, website)

    new_page()
    sy = PAGE_H - CV_HEADER_H - 28
    sidebar_x = 12
    sidebar_w = CV_SIDEBAR_W - 20

    for section_title, items in [
        ("Expertise", pdf_data.get("expertise") or []),
        ("Tools", pdf_data.get("tools") or []),
    ]:
        if items:
            sy = section_label(c, section_title, sidebar_x, sy, sidebar_w)
            for item in items:
                sy = wrap(c, "\u00b7 " + item, sidebar_x, sy, sidebar_w, F("Sans"), 7.5, SLATE, lh=11)
            sy -= 10

    for edu in (pdf_data.get("education") or []):
        sy = section_label(c, "Education", sidebar_x, sy, sidebar_w)
        c.setFillColor(NIGHT)
        c.setFont(F("Sans-Bold"), 7.5)
        c.drawString(sidebar_x, sy, edu.get("degree") or "")
        sy -= 11
        c.setFillColor(MID)
        c.setFont(F("Sans"), 7)
        c.drawString(sidebar_x, sy, edu.get("school") or "")
        sy -= 10
        c.drawString(sidebar_x, sy, edu.get("years") or "")
        sy -= 14

    for cert in (pdf_data.get("certifications") or []):
        sy = section_label(c, "Certifications", sidebar_x, sy, sidebar_w)
        sy = wrap(c, cert.get("name") or "", sidebar_x, sy, sidebar_w, F("Sans-Bold"), 7.5, NIGHT, lh=11)
        c.setFillColor(MID)
        c.setFont(F("Sans"), 7)
        c.drawString(sidebar_x, sy, cert.get("org") or "")
        sy -= 14

    cy = PAGE_H - CV_HEADER_H - 24
    if summary_text:
        cy = section_label(c, "Profile", CV_CONTENT_X, cy, CV_CONTENT_W)
        cy = wrap(c, summary_text, CV_CONTENT_X, cy, CV_CONTENT_W, F("Serif-Italic"), 8.8, SLATE, lh=14)
        cy -= 14

    cy = section_label(c, "Experience", CV_CONTENT_X, cy, CV_CONTENT_W)
    for exp in (pdf_data.get("experienceItems") or []):
        if cy < 100:
            c.showPage(); new_page(); cy = PAGE_H - CV_HEADER_H - 24
        c.setFillColor(NIGHT)
        c.setFont(F("Sans-Bold"), 9.5)
        c.drawString(CV_CONTENT_X, cy, exp.get("role") or "")
        cy -= 13
        c.setFillColor(GOLD)
        c.setFont(F("Sans-Med"), 8)
        company_line = " \u00b7 ".join(filter(None, [exp.get("company"), exp.get("location")]))
        c.drawString(CV_CONTENT_X, cy, company_line)
        c.setFillColor(MID)
        c.setFont(F("Sans"), 7.5)
        c.drawRightString(CV_CONTENT_X + CV_CONTENT_W, cy, exp.get("dates") or "")
        cy -= 11
        for b in (exp.get("bullets") or []):
            cy = bullet_line(c, b, CV_CONTENT_X + 2, cy, CV_CONTENT_W)
        cy -= 10

    c.showPage(); c.save()
    return buf.getvalue()


# --------------------------------------------------------------------------- #
#  PDF — TEMPLATE B (clean single column)                                     #
# --------------------------------------------------------------------------- #
TB_HEADER_H = 72
TB_MARGIN   = 44
TB_CONTENT_W = PAGE_W - TB_MARGIN * 2


def cv_header_b(c, name, title_line, contact_items):
    c.setFillColor(NIGHT)
    c.rect(0, PAGE_H - TB_HEADER_H, PAGE_W, TB_HEADER_H, fill=1, stroke=0)
    c.setFillColor(GOLD)
    c.rect(0, PAGE_H - 4, PAGE_W, 4, fill=1, stroke=0)
    c.setFillColor(WHITE)
    c.setFont(F("Serif"), 20)
    c.drawString(TB_MARGIN, PAGE_H - 32, name)
    if title_line:
        c.setFillColor(GOLD)
        c.setFont(F("Sans-Med"), 7.5)
        c.drawString(TB_MARGIN, PAGE_H - 46, title_line)
    contact_str = "  \u00b7  ".join(contact_items[:5])
    c.setFillColor(WHITE)
    c.setFont(F("Sans"), 7)
    c.drawString(TB_MARGIN, PAGE_H - 62, contact_str)


def build_cv_pdf_b(cv: dict, pdf_data: dict) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)

    name = cv.get("candidateName") or "Candidate"
    email = cv.get("candidateEmail") or ""
    phone = cv.get("candidatePhone") or ""
    linkedin = cv.get("candidateLinkedIn") or ""
    website = cv.get("candidateWebsite") or ""
    title_line = cv.get("titleLine") or ""
    location = cv.get("candidateLocation") or ""
    summary_text = cv.get("summary") or ""
    contact_items = [i for i in [email, phone, linkedin, location] if i]

    def new_page():
        cv_header_b(c, name, title_line, contact_items)
        footer(c, name, email, website)

    new_page()
    cy = PAGE_H - TB_HEADER_H - 28

    if summary_text:
        cy = section_label(c, "Profile", TB_MARGIN, cy, TB_CONTENT_W)
        cy = wrap(c, summary_text, TB_MARGIN, cy, TB_CONTENT_W, F("Serif-Italic"), 9, SLATE, lh=14)
        cy -= 14
        hr(c, TB_MARGIN, cy, TB_MARGIN + TB_CONTENT_W, color=GOLD, lw=0.4)
        cy -= 14

    all_skills = list(pdf_data.get("expertise") or []) + list(pdf_data.get("tools") or [])
    if all_skills:
        cy = section_label(c, "Skills & Tools", TB_MARGIN, cy, TB_CONTENT_W)
        cy = wrap(c, "  \u00b7  ".join(all_skills), TB_MARGIN, cy, TB_CONTENT_W, F("Sans"), 8, SLATE, lh=12)
        cy -= 14
        hr(c, TB_MARGIN, cy, TB_MARGIN + TB_CONTENT_W, color=GOLD, lw=0.4)
        cy -= 14

    cy = section_label(c, "Experience", TB_MARGIN, cy, TB_CONTENT_W)
    for exp in (pdf_data.get("experienceItems") or []):
        if cy < 100:
            c.showPage(); new_page(); cy = PAGE_H - TB_HEADER_H - 28
        c.setFillColor(NIGHT)
        c.setFont(F("Sans-Bold"), 9.5)
        c.drawString(TB_MARGIN, cy, exp.get("role") or "")
        c.setFillColor(MID)
        c.setFont(F("Sans"), 7.5)
        c.drawRightString(TB_MARGIN + TB_CONTENT_W, cy, exp.get("dates") or "")
        cy -= 13
        c.setFillColor(GOLD)
        c.setFont(F("Sans-Med"), 8)
        company_line = " \u00b7 ".join(filter(None, [exp.get("company"), exp.get("location")]))
        c.drawString(TB_MARGIN, cy, company_line)
        cy -= 13
        for b in (exp.get("bullets") or []):
            cy = bullet_line(c, b, TB_MARGIN + 2, cy, TB_CONTENT_W)
        cy -= 8
        hr(c, TB_MARGIN, cy, TB_MARGIN + TB_CONTENT_W, color=LIGHT, lw=0.3)
        cy -= 10

    edu_list = pdf_data.get("education") or []
    if edu_list:
        cy = section_label(c, "Education", TB_MARGIN, cy, TB_CONTENT_W)
        for edu in edu_list:
            c.setFillColor(NIGHT)
            c.setFont(F("Sans-Bold"), 8.5)
            c.drawString(TB_MARGIN, cy, edu.get("degree") or "")
            cy -= 12
            c.setFillColor(MID)
            c.setFont(F("Sans"), 8)
            c.drawString(TB_MARGIN, cy, f"{edu.get('school', '')}  \u00b7  {edu.get('years', '')}")
            cy -= 14

    c.showPage(); c.save()
    return buf.getvalue()


# --------------------------------------------------------------------------- #
#  PDF — COVER LETTER                                                         #
# --------------------------------------------------------------------------- #
CL_HEADER_H = 70
CL_MARGIN   = 50
CL_TEXT_W   = PAGE_W - CL_MARGIN * 2


def letter_header(c, name, title_line, contact_items):
    c.setFillColor(NIGHT)
    c.rect(0, PAGE_H - CL_HEADER_H, PAGE_W, CL_HEADER_H, fill=1, stroke=0)
    c.setFillColor(GOLD)
    c.rect(0, PAGE_H - 4, PAGE_W, 4, fill=1, stroke=0)
    c.setFillColor(WHITE)
    c.setFont(F("Serif"), 18)
    c.drawString(CL_MARGIN, PAGE_H - 30, name)
    if title_line:
        c.setFillColor(GOLD)
        c.setFont(F("Sans-Med"), 7.5)
        c.drawString(CL_MARGIN, PAGE_H - 44, title_line)
    contact_str = "  \u00b7  ".join(contact_items[:4])
    c.setFillColor(WHITE)
    c.setFont(F("Sans"), 7)
    c.drawString(CL_MARGIN, PAGE_H - 58, contact_str)


def build_cover_letter_pdf(cv: dict, letter: dict, date_str: str, salutation: str) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    name = cv.get("candidateName") or "Candidate"
    email = cv.get("candidateEmail") or ""
    website = cv.get("candidateWebsite") or ""
    title_line = cv.get("titleLine") or ""
    contact_items = [i for i in [email, cv.get("candidatePhone"), cv.get("candidateLinkedIn")] if i]
    tagline = letter.get("tagline") or ""
    signoff_name = letter.get("signoffName") or name

    paras = [p for p in [
        letter.get("opening"), letter.get("body1"),
        letter.get("body2"), letter.get("body3Remote"), letter.get("closing"),
    ] if p]

    def new_page():
        letter_header(c, name, title_line, contact_items)
        footer(c, name, email, website)

    new_page()
    cy = PAGE_H - CL_HEADER_H - 44
    c.setFillColor(MID)
    c.setFont(F("Sans"), 8)
    c.drawString(CL_MARGIN, cy, date_str)
    cy -= 28
    c.setFillColor(NIGHT)
    c.setFont(F("Serif"), 10)
    c.drawString(CL_MARGIN, cy, salutation)
    cy -= 24

    for p in paras:
        if cy < 110:
            c.showPage(); new_page(); cy = PAGE_H - CL_HEADER_H - 44
        cy = wrap(c, p, CL_MARGIN, cy, CL_TEXT_W, F("Serif"), 9.8, SLATE, lh=16)
        cy -= 18

    if cy < 120:
        c.showPage(); new_page(); cy = PAGE_H - CL_HEADER_H - 44

    c.setFillColor(NIGHT)
    c.setFont(F("Serif"), 10)
    c.drawString(CL_MARGIN, cy, "Warm regards,")
    cy -= 28
    c.setFont(F("Serif"), 13)
    c.drawString(CL_MARGIN, cy, signoff_name)
    cy -= 8
    hr(c, CL_MARGIN, cy, CL_MARGIN + 180, color=GOLD, lw=0.8)
    cy -= 12
    if tagline:
        c.setFillColor(MID)
        c.setFont(F("Sans-Light"), 7.5)
        c.drawString(CL_MARGIN, cy, tagline)

    c.showPage(); c.save()
    return buf.getvalue()


# --------------------------------------------------------------------------- #
#  PDF — INTERVIEW PREP                                                       #
# --------------------------------------------------------------------------- #
def build_prep_pdf(cv: dict, prep: dict) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    name = cv.get("candidateName") or "Candidate"
    email = cv.get("candidateEmail") or ""
    website = cv.get("candidateWebsite") or ""
    title_line = cv.get("titleLine") or ""
    contact_items = [i for i in [email, cv.get("candidatePhone"), cv.get("candidateLinkedIn")] if i]
    questions = prep.get("questions") or []

    def new_page():
        letter_header(c, name, title_line, contact_items)
        footer(c, name, email, website)

    new_page()
    cy = PAGE_H - CL_HEADER_H - 40
    cy = section_label(c, "Interview Prep Pack", CL_MARGIN, cy, CL_TEXT_W)
    cy -= 6

    def ensure_space(needed):
        nonlocal cy
        if cy - needed < 60:
            c.showPage(); new_page(); cy = PAGE_H - CL_HEADER_H - 40

    for idx, q in enumerate(questions, start=1):
        ensure_space(90)
        c.setFillColor(NIGHT)
        c.setFont(F("Sans-Bold"), 10)
        for line in _wrap_lines(c, f"{idx}.  {q.get('question', '')}", CL_TEXT_W, F("Sans-Bold"), 10):
            c.drawString(CL_MARGIN, cy, line); cy -= 15
        cy -= 2
        if q.get("weakness"):
            c.setFillColor(GOLD)
            c.setFont(F("Sans-Italic"), 8)
            for line in _wrap_lines(c, "Targets: " + q["weakness"], CL_TEXT_W, F("Sans-Italic"), 8):
                c.drawString(CL_MARGIN, cy, line); cy -= 12
        cy -= 6
        star = q.get("starAnswer") or {}
        for label, key in (("Situation", "situation"), ("Task", "task"),
                           ("Action", "action"), ("Result", "result")):
            body = star.get(key, "")
            if not body:
                continue
            ensure_space(40)
            c.setFillColor(GOLD)
            c.setFont(F("Sans-Bold"), 7)
            c.drawString(CL_MARGIN, cy, label.upper()); cy -= 12
            cy = wrap(c, body, CL_MARGIN + 4, cy, CL_TEXT_W - 4, F("Sans"), 8.5, SLATE, lh=13)
            cy -= 6
        if q.get("needsRealExample"):
            ensure_space(34)
            box_h = 22
            c.setStrokeColor(GOLD); c.setLineWidth(0.8)
            c.rect(CL_MARGIN, cy - box_h + 6, CL_TEXT_W, box_h, fill=0, stroke=1)
            c.setFillColor(GOLD); c.setFont(F("Sans-Bold"), 7.5)
            c.drawString(CL_MARGIN + 8, cy - 6,
                         "[NEEDS REAL EXAMPLE]  Insert a concrete achievement before the interview.")
            cy -= box_h + 8
        cy -= 20

    c.showPage(); c.save()
    return buf.getvalue()


# --------------------------------------------------------------------------- #
#  DOCX BUILDERS                                                              #
# --------------------------------------------------------------------------- #
def _set_gold(run): run.font.color.rgb = RGBColor(0xB8, 0x9A, 0x68)
def _set_dark(run): run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
def _set_mid(run):  run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def build_cv_docx(cv: dict, pdf_data: dict, template: str = "A") -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6); section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8); section.right_margin = Inches(0.8)

    name = cv.get("candidateName") or "Candidate"
    title_line = cv.get("titleLine") or ""
    email = cv.get("candidateEmail") or ""
    phone = cv.get("candidatePhone") or ""
    linkedin = cv.get("candidateLinkedIn") or ""
    location = cv.get("candidateLocation") or ""
    summary_text = cv.get("summary") or ""

    h = doc.add_paragraph()
    run = h.add_run(name); run.bold = True; run.font.size = Pt(22); _set_dark(run)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if title_line:
        p = doc.add_paragraph(); run = p.add_run(title_line); run.font.size = Pt(9); _set_gold(run)

    contact_parts = [x for x in [email, phone, linkedin, location] if x]
    if contact_parts:
        p = doc.add_paragraph(); run = p.add_run("  \u00b7  ".join(contact_parts))
        run.font.size = Pt(8); _set_mid(run)

    doc.add_paragraph()

    def add_section_heading(title):
        p = doc.add_paragraph(); run = p.add_run(title.upper())
        run.bold = True; run.font.size = Pt(7); _set_gold(run)
        p.paragraph_format.space_after = Pt(2)

    if summary_text:
        add_section_heading("Profile")
        p = doc.add_paragraph(summary_text); p.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    all_skills = list(pdf_data.get("expertise") or []) + list(pdf_data.get("tools") or [])
    if all_skills:
        add_section_heading("Skills & Tools")
        p = doc.add_paragraph("  \u00b7  ".join(all_skills)); p.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    if pdf_data.get("experienceItems"):
        add_section_heading("Experience")
        for exp in pdf_data["experienceItems"]:
            p = doc.add_paragraph()
            r1 = p.add_run(exp.get("role") or ""); r1.bold = True; r1.font.size = Pt(10); _set_dark(r1)
            p2 = doc.add_paragraph()
            r2 = p2.add_run(" \u00b7 ".join(filter(None, [exp.get("company"), exp.get("location")])))
            r2.font.size = Pt(8.5); _set_gold(r2)
            r3 = p2.add_run(f"  {exp.get('dates', '')}"); r3.font.size = Pt(8); _set_mid(r3)
            for b in (exp.get("bullets") or []):
                p3 = doc.add_paragraph(style="List Bullet")
                run = p3.add_run(b); run.font.size = Pt(9)
            doc.add_paragraph()

    if pdf_data.get("education"):
        add_section_heading("Education")
        for edu in pdf_data["education"]:
            p = doc.add_paragraph(); r = p.add_run(edu.get("degree") or ""); r.bold = True; r.font.size = Pt(9)
            p2 = doc.add_paragraph(); r2 = p2.add_run(f"{edu.get('school', '')}  \u00b7  {edu.get('years', '')}")
            r2.font.size = Pt(8.5); _set_mid(r2)
        doc.add_paragraph()

    if pdf_data.get("certifications"):
        add_section_heading("Certifications")
        for cert in pdf_data["certifications"]:
            p = doc.add_paragraph()
            r = p.add_run(cert.get("name") or ""); r.bold = True; r.font.size = Pt(9)
            r2 = p.add_run(f"  \u2014  {cert.get('org', '')}"); r2.font.size = Pt(8.5); _set_mid(r2)

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def build_cover_letter_docx(cv: dict, letter: dict, date_str: str, salutation: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8); section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0); section.right_margin = Inches(1.0)

    name = cv.get("candidateName") or "Candidate"
    title_line = cv.get("titleLine") or ""

    h = doc.add_paragraph(); r = h.add_run(name); r.bold = True; r.font.size = Pt(18)
    if title_line:
        p = doc.add_paragraph(); r = p.add_run(title_line); r.font.size = Pt(9); _set_gold(r)
    doc.add_paragraph()

    if date_str:
        p = doc.add_paragraph(date_str); p.runs[0].font.size = Pt(9); _set_mid(p.runs[0])
        doc.add_paragraph()
    if salutation:
        p = doc.add_paragraph(salutation); p.runs[0].font.size = Pt(10); doc.add_paragraph()

    paras = [letter.get(k) for k in ["opening", "body1", "body2", "body3Remote", "closing"] if letter.get(k)]
    for para in paras:
        p = doc.add_paragraph(para); p.runs[0].font.size = Pt(10); p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()
    p = doc.add_paragraph("Warm regards,"); p.runs[0].font.size = Pt(10)
    p2 = doc.add_paragraph(letter.get("signoffName") or name)
    r = p2.runs[0]; r.bold = True; r.font.size = Pt(12)
    if letter.get("tagline"):
        p3 = doc.add_paragraph(letter["tagline"]); p3.runs[0].font.size = Pt(8); _set_mid(p3.runs[0])

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def build_prep_docx(cv: dict, prep: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7); section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9); section.right_margin = Inches(0.9)

    h = doc.add_paragraph()
    r = h.add_run(f"Interview Prep Pack \u2014 {cv.get('candidateName', 'Candidate')}")
    r.bold = True; r.font.size = Pt(14)
    doc.add_paragraph()

    for idx, q in enumerate((prep.get("questions") or []), start=1):
        p = doc.add_paragraph(); r = p.add_run(f"{idx}. {q.get('question', '')}")
        r.bold = True; r.font.size = Pt(11)
        if q.get("weakness"):
            p2 = doc.add_paragraph(f"Targets: {q['weakness']}")
            p2.runs[0].font.size = Pt(9); _set_gold(p2.runs[0])
        star = q.get("starAnswer") or {}
        for label, key in (("Situation", "situation"), ("Task", "task"),
                           ("Action", "action"), ("Result", "result")):
            body = star.get(key, "")
            if body:
                p3 = doc.add_paragraph()
                r1 = p3.add_run(f"{label.upper()}: "); r1.bold = True; r1.font.size = Pt(9); _set_gold(r1)
                r2 = p3.add_run(body); r2.font.size = Pt(9)
        if q.get("needsRealExample"):
            p4 = doc.add_paragraph("[NEEDS REAL EXAMPLE] \u2014 Add a concrete achievement before the interview.")
            p4.runs[0].font.size = Pt(9); _set_gold(p4.runs[0])
        doc.add_paragraph()

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
#  HELPERS                                                                    #
# --------------------------------------------------------------------------- #
def clean_company(name):
    if not name:
        return "Company"
    name = re.sub(r"\b(Ltd|Inc|LLC|Limited|Incorporated|Corp|GmbH|PLC|Co)\b\.?",
                  "", name, flags=re.IGNORECASE)
    name = re.sub(r"[^\w\s-]", "", name).strip()
    return re.sub(r"\s+", "_", name) or "Company"


def split_name(full):
    parts = (full or "Candidate").strip().split()
    return (parts[0], parts[-1]) if len(parts) > 1 else (parts[0], "CV")


# --------------------------------------------------------------------------- #
#  ROUTE — GENERATE PDFs                                                      #
# --------------------------------------------------------------------------- #
@app.route("/generate-pdfs", methods=["POST"])
def generate_pdfs():
    try:
        data = request.get_json(force=True) or {}
        cv = data.get("rewrittenCV") or {}
        pdf_data = data.get("pdfData") or {}
        letter = data.get("coverLetter") or {}
        prep = data.get("interviewPrep") or {}
        company = data.get("companyName") or cv.get("companyName") or "Company"
        date_str = data.get("dateStr") or ""
        salutation = data.get("salutation") or "Dear Hiring Manager,"
        template = data.get("template") or "A"

        first, last = split_name(cv.get("candidateName"))
        comp = clean_company(company)
        cv_bytes = build_cv_pdf_a(cv, pdf_data) if template != "B" else build_cv_pdf_b(cv, pdf_data)

        files = [{
            "name": f"{first}_{last}_CV_{comp}.pdf", "kind": "cv", "format": "pdf",
            "description": "Rewritten, ATS-safe CV — branded PDF.",
            "data": base64.b64encode(cv_bytes).decode(),
        }]
        if letter:
            cover_bytes = build_cover_letter_pdf(cv, letter, date_str, salutation)
            files.append({"name": f"{first}_{last}_CoverLetter_{comp}.pdf", "kind": "cover",
                          "format": "pdf", "description": "Tailored cover letter.",
                          "data": base64.b64encode(cover_bytes).decode()})
        if prep:
            prep_bytes = build_prep_pdf(cv, prep)
            files.append({"name": f"{first}_{last}_InterviewPrep_{comp}.pdf", "kind": "prep",
                          "format": "pdf", "description": "Targeted STAR questions.",
                          "data": base64.b64encode(prep_bytes).decode()})
        return jsonify({"status": "success", "files": files})
    except Exception as e:
        return fail("Could not generate PDF documents.", 500, exc=e)


# --------------------------------------------------------------------------- #
#  ROUTE — GENERATE DOCX                                                      #
# --------------------------------------------------------------------------- #
@app.route("/generate-docx", methods=["POST"])
def generate_docx():
    try:
        data = request.get_json(force=True) or {}
        cv = data.get("rewrittenCV") or {}
        pdf_data = data.get("pdfData") or {}
        letter = data.get("coverLetter") or {}
        prep = data.get("interviewPrep") or {}
        company = data.get("companyName") or cv.get("companyName") or "Company"
        date_str = data.get("dateStr") or ""
        salutation = data.get("salutation") or "Dear Hiring Manager,"
        template = data.get("template") or "A"

        first, last = split_name(cv.get("candidateName"))
        comp = clean_company(company)
        cv_bytes = build_cv_docx(cv, pdf_data, template)

        files = [{
            "name": f"{first}_{last}_CV_{comp}.docx", "kind": "cv", "format": "docx",
            "description": "Rewritten CV — editable Word document.",
            "data": base64.b64encode(cv_bytes).decode(),
        }]
        if letter:
            cover_bytes = build_cover_letter_docx(cv, letter, date_str, salutation)
            files.append({"name": f"{first}_{last}_CoverLetter_{comp}.docx", "kind": "cover",
                          "format": "docx", "description": "Cover letter — editable Word document.",
                          "data": base64.b64encode(cover_bytes).decode()})
        if prep:
            prep_bytes = build_prep_docx(cv, prep)
            files.append({"name": f"{first}_{last}_InterviewPrep_{comp}.docx", "kind": "prep",
                          "format": "docx", "description": "Interview prep pack — editable Word document.",
                          "data": base64.b64encode(prep_bytes).decode()})
        return jsonify({"status": "success", "files": files})
    except Exception as e:
        return fail("Could not generate Word documents.", 500, exc=e)


# --------------------------------------------------------------------------- #
#  HEALTH                                                                     #
# --------------------------------------------------------------------------- #
@app.route("/health", methods=["GET"])
def health():
    return jsonify({
        "status": "ok",
        "service": "Sofia",
        "aiConfigured": _anthropic is not None,
        "creditsEnforced": ENFORCE_CREDITS,
        "fonts": _REGISTERED,
    })


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=DEBUG_MODE)
